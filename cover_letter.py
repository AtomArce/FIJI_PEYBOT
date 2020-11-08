import re
from docx import Document
from docx2pdf import convert

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)



regex1 = re.compile(r"job_name")
regex2 = re.compile(r"company_name")
replace1 = r"Intern"
replace2 = r"Google"
filename = "coverletter.docx"
doc = Document(filename)
docx_replace_regex(doc, regex1, replace1)
docx_replace_regex(doc, regex2, replace2)
doc.save('new.docx')

convert("new.docx")
convert("new.docx", "new.pdf")
